import openai
import docx
from tqdm import tqdm
from docx.oxml import OxmlElement
import time

# Set up OpenAI API - You must provide your own API Key here:
openai.api_key = 'YOURAPIKEY'


def modernize_text(text, max_retries=3, retry_delay=5):
    for attempt in range(max_retries):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-1106-preview",
                # You can modify this prompt to your specific use case. In this example we were preparing to translate an old English book and so we used ChatGPT to modernize it as a help to the translator.
                messages=[{"role": "system", "content": "You are a helpful assistant that modernizes sentences from old English books into modern English. Don't use commas too much and change to active voice when possible and if it is just a name, leave it as just the name. I will give you a sentence and you will modernize it without any comment."}, 
                          {"role": "user", "content": text}]
            )
            return response.choices[0].message['content']
        except openai.error.Timeout as e:
            if attempt < max_retries - 1:
                time.sleep(retry_delay)
            else:
                raise e  # Re-raise the exception on the last attempt


def set_cell_borders(cell, border_width_pt=1):
    # This function adds borders to the table in the exported Word document
    # Set borders for a table cell
    cell_properties = cell._element.get_or_add_tcPr()

    # Specify the border types to be added
    borders = ['top', 'bottom', 'left', 'right']

    for border in borders:
        tag = 'w:{}'.format(border)
        border_element = OxmlElement(tag)
        border_element.set(docx.oxml.ns.qn('w:val'), 'single')
        border_element.set(docx.oxml.ns.qn('w:sz'), str(border_width_pt * 8))  # Word's measurement unit is 1/8th of a point
        border_element.set(docx.oxml.ns.qn('w:space'), '0')
        border_element.set(docx.oxml.ns.qn('w:color'), 'auto')
        cell_properties.append(border_element)

def process_document(input_path, output_path):
    # This function processes the input Microsoft Word Document and sends it, paragraph by paragraphy, to ChatGPT and then exports a Word Document with both the original and modernized versions side by side.
    doc = docx.Document(input_path)
    new_doc = docx.Document()

    # Wrap the loop with tqdm for the progress bar
    for paragraph in tqdm(doc.paragraphs, desc="Processing", unit="para"):
        original_text = paragraph.text

        # Skip the paragraph if it is blank or contains only spaces
        if original_text.strip() == '':
            continue

        modernized_text = modernize_text(original_text)

        # Create a table for the two-column layout
        table = new_doc.add_table(rows=1, cols=2)

        # Apply border to each cell
        for cell in table.columns[0].cells + table.columns[1].cells:
            set_cell_borders(cell)

        # Get the row
        row = table.rows[0]

        # Set the original text in the first column
        row.cells[0].text = original_text
        # Set the modernized text in the second column
        row.cells[1].text = modernized_text

        # Add a blank line after each table for spacing
        new_doc.add_paragraph()

    new_doc.save(output_path)

# This gets called when you run with Python via terminal: python3 translation-helper.py
process_document('input.docx', 'output.docx')
